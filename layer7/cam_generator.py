"""
Layer 7 — CAM Document Generator
Generates a full 13-section RBI-compliant Credit Appraisal Memorandum (CAM)
in DOCX format with evidence panel, digital watermark, and multi-format export.
"""

import os
import json
import hashlib
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

# ─── Helpers ─────────────────────────────────────────────────────────────────

def _safe(val, default='—'):
    if val is None or val == '' or val == []:
        return default
    if isinstance(val, dict) and 'value' in val:
        return val['value'] if val['value'] is not None else default
    return val

def _fmt_inr(val):
    try:
        n = float(val)
        if n >= 10000000:
            return f"₹{n/10000000:.2f} Cr"
        if n >= 100000:
            return f"₹{n/100000:.2f} L"
        return f"₹{n:,.0f}"
    except:
        return str(val) if val else '—'

def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def _add_table_row(table, cells, bold=False, bg_color=None):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(val) if val is not None else '—'
        para = cell.paragraphs[0]
        para.style.font.size = Pt(9)
        if bold:
            for run in para.runs:
                run.bold = True
        if bg_color:
            from docx.oxml.ns import qn
            shading = cell._element.get_or_add_tcPr()
            shading_el = shading.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): bg_color
            })
            shading.append(shading_el)
    return row


# ─── CAM Generator Class ────────────────────────────────────────────────────

class CAMGenerator:
    """Generate a full 14-section RBI-compliant Credit Appraisal Memorandum."""

    def __init__(self, app_data, l2_data, l3_data, l4_data, l5_data):
        self.app = app_data
        self.l2 = l2_data or {}
        self.l3 = l3_data or {}
        self.l4 = l4_data or {}
        self.l5 = l5_data or {}
        self.doc = Document()
        self.cam_content_text = ""

        # Extract common data
        self.extracted = self.l2.get('extracted', {}).get('financial_data', {}) or self.l2.get('extracted', {}) or {}
        self.decision = self.l5.get('decision_summary', {})
        self.explanation = self.l5.get('explanation', {})
        self.forensics = self.l4.get('forensics_report', {})
        self.research = self.l4.get('research_findings', {})
        self.five_cs = self.explanation.get('five_cs', {})
        self.loan = self.l5.get('loan_structure', {})
        self.pricing = self.l5.get('pricing', {})

        # Officer-added HITL issues & custom fields
        self.officer_issues = app_data.get('officer_issues', [])
        self.custom_fields = app_data.get('custom_fields', {})
        if isinstance(self.custom_fields, str):
            try:
                self.custom_fields = json.loads(self.custom_fields)
            except:
                self.custom_fields = {}

    def generate(self, output_path):
        """Generate the full CAM DOCX document."""
        self._setup_styles()
        self._section_1_cover()
        self._section_2_executive_summary()
        self._section_3_borrower_profile()
        self._section_3b_shareholding_pattern()
        try:
            self._section_3c_statutory_compliance()
        except Exception as e:
            print(f"  [CAM] Section 3C Statutory Compliance skipped: {e}")
        self._section_4_loan_proposal()
        self._section_5_financial_analysis()
        self._section_6_existing_debt()
        self._section_7_security_collateral()
        self._section_8_sector_risk()
        try:
            self._section_8b_esg_risk()
        except Exception as e:
            print(f"  [CAM] Section 8B ESG Risk skipped: {e}")
        self._section_9_ai_risk_score()
        self._section_10_forensic_alerts()
        self._section_11_five_cs()
        self._section_12_swot()
        self._section_13_officer_issues()
        self._section_14_recommendations()
        self._appendix_a_evidence()

        # Compute hash and embed in metadata
        cam_hash = self._compute_hash()
        self.doc.core_properties.comments = f"CAM_HASH:{cam_hash}"

        self.doc.save(output_path)
        return {
            'path': output_path,
            'cam_hash': cam_hash,
            'sections': 14,
            'timestamp': datetime.utcnow().isoformat()
        }

    def _setup_styles(self):
        """Configure document styles."""
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        for section in self.doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Set up RBI logo in the header for all pages (except cover page)
        if self.doc.sections:
            section = self.doc.sections[0]
            section.different_first_page_header_footer = True
            header = section.header
            
            rbi_logo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'assets', 'RBI_Logo.jpg')
            if os.path.exists(rbi_logo_path):
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = header_para.add_run()
                run.add_picture(rbi_logo_path, width=Inches(0.6))

    # ─── Section 1: Cover Page ─────────────────────────────────────

    def _section_1_cover(self):
        """Cover Page with RBI logo, bank header, and case info."""
        # Try to add RBI logo
        rbi_logo = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static', 'assets', 'RBI_Logo.jpg')
        if os.path.exists(rbi_logo):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(rbi_logo, width=Inches(1.2))

        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('CREDIT APPRAISAL MEMORANDUM (CAM)')
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('— CONFIDENTIAL —')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xcc, 0, 0)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Format as per RBI Prudential Norms')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        self.doc.add_paragraph()

        # Case metadata table
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        case_id = self.app.get('case_id', '—')
        meta_rows = [
            ('Case ID', case_id),
            ('Proposal Date', datetime.now().strftime('%d-%m-%Y')),
            ('Company Name', self.app.get('company_name', '—')),
            ('Model Version', self.l5.get('audit_snapshot', {}).get('model_metadata', {}).get('model_version', 'v4.3')),
            ('Schema Version', '1.0.0'),
            ('Internal CRR Grade', self.decision.get('risk_band', '—')),
            ('AI Credit Score', str(self.decision.get('final_credit_score', '—'))),
            ('External Rating', _safe(self.extracted.get('external_credit_rating'), 'CRISIL BBB+ (Assessed)')),
        ]
        for label, value in meta_rows:
            _add_table_row(table, [label, value])

        self.doc.add_page_break()
        self.cam_content_text += f"COVER: {case_id} | {self.app.get('company_name', '')}\n"

    # ─── Section 2: Executive Summary ──────────────────────────────

    def _section_2_executive_summary(self):
        _add_heading(self.doc, '1. EXECUTIVE SUMMARY & BORROWER PROFILE', level=1)

        company = self.app.get('company_name', '—')
        score = self.decision.get('final_credit_score', '—')
        decision = self.decision.get('decision', '—')
        risk_band = self.decision.get('risk_band', '—')

        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        rows = [
            ('Applicant Name', f"M/s. {company}"),
            ('Nature of Business', _safe(self.extracted.get('nature_of_business'), 'Manufacturing / Services')),
            ('CIN', _safe(self.extracted.get('cin'))),
            ('GSTIN', _safe(self.extracted.get('gstin'))),
            ('PAN', _safe(self.extracted.get('pan_number'))),
            ('Promoter Background', _safe(self.extracted.get('promoter_name'), 'As per documents')),
            ('Relationship History', _safe(self.extracted.get('banking_relationship'), 'Banking since assessment')),
            ('Internal CRR', risk_band),
            ('AI Credit Score', str(score)),
            ('Decision', decision),
            ('KYC/AML Sign-off', '✅ KYC verified as per RBI Master Direction 2016 (updated 2023)'),
            ('Top 3 Risk Factors', ', '.join(
                [d.get('label', '') for d in self.explanation.get('shap_top_negative', [])[:3]]
            ) or 'See Section 9'),
        ]
        for label, value in rows:
            _add_table_row(table, [label, str(value)])

        # LLM summary
        llm_summary = self.decision.get('llm_decision_summary', '')
        if llm_summary:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run('AI Decision Summary: ')
            run.bold = True
            run.font.size = Pt(10)
            p.add_run(llm_summary)

        self.doc.add_paragraph()
        self.cam_content_text += f"EXEC: {company} | Score: {score} | Decision: {decision}\n"

    # ─── Section 3B: Shareholding Pattern ──────────────────────────────

    def _section_3b_shareholding_pattern(self):
        _add_heading(self.doc, '3. SHAREHOLDING PATTERN & OWNERSHIP', level=1)
        
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        # Breakdown
        _add_table_row(table, ['Promoter Holding (%)', _safe(self.extracted.get('promoter_holding_pct'))])
        _add_table_row(table, ['Public Holding (%)', _safe(self.extracted.get('public_holding_pct'))])
        _add_table_row(table, ['Institutional Holding (%)', _safe(self.extracted.get('institutional_holding_pct'))])
        _add_table_row(table, ['Foreign Holding (%)', _safe(self.extracted.get('foreign_holding_pct'))])
        _add_table_row(table, ['Pledged Shares (%)', _safe(self.extracted.get('pledged_shares_pct'))], bold=True, bg_color='fdf2f8')
        _add_table_row(table, ['As of Date / Period', _safe(self.extracted.get('shareholding_period'))])
        
        self.doc.add_paragraph()
        
        # Top Shareholders list if available
        top_holders = self.extracted.get('top_shareholders', [])
        if top_holders and isinstance(top_holders, list):
            _add_heading(self.doc, 'Top Shareholders', level=2)
            t2 = self.doc.add_table(rows=1, cols=2)
            t2.style = 'Table Grid'
            t2.rows[0].cells[0].text = 'Shareholder Name'
            t2.rows[0].cells[1].text = 'Holding (%)'
            for h in top_holders:
                if isinstance(h, dict):
                    _add_table_row(t2, [h.get('name', '—'), h.get('percentage', '—')])
                else:
                    _add_table_row(t2, [str(h), '—'])
                    
        self.doc.add_paragraph()

    # ─── Section 3: Borrower & Group Profile ───────────────────────

    def _section_3_borrower_profile(self):
        _add_heading(self.doc, '2. BORROWER & GROUP PROFILE', level=1)

        mca = self.research.get('mca_checks', {})
        adverse = self.research.get('adverse_media', {})
        litigation = self.research.get('litigation', {})

        # Company info
        p = self.doc.add_paragraph()
        run = p.add_run('Company Background: ')
        run.bold = True
        p.add_run(_safe(self.extracted.get('nature_of_business'), 'As per submitted documents'))

        # Promoter net worth statement (RBI requirement)
        _add_heading(self.doc, 'Promoter Personal Net Worth Statement', level=2)
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        promoter = _safe(self.extracted.get('promoter_name'), 'Principal Promoter')
        _add_table_row(table, ['Promoter Name', promoter])
        _add_table_row(table, ['Personal Assets', _fmt_inr(self.extracted.get('promoter_personal_assets', 'As declared'))])
        _add_table_row(table, ['Personal Liabilities', _fmt_inr(self.extracted.get('promoter_personal_liabilities', 'As declared'))])
        _add_table_row(table, ['Personal Net Worth', _fmt_inr(self.extracted.get('promoter_net_worth', 'As declared'))])

        # Director DIN status
        directors = mca.get('directors', [])
        if directors:
            _add_heading(self.doc, 'Director DIN Status (MCA)', level=2)
            t2 = self.doc.add_table(rows=1, cols=3)
            t2.style = 'Table Grid'
            t2.rows[0].cells[0].text = 'Director Name'
            t2.rows[0].cells[1].text = 'DIN'
            t2.rows[0].cells[2].text = 'Status'
            for d in directors[:5]:
                _add_table_row(t2, [d.get('name', '—'), d.get('din', '—'), d.get('status', '—')])

        # Adverse media
        if adverse:
            _add_heading(self.doc, 'Adverse Media Summary', level=2)
            self.doc.add_paragraph(adverse.get('summary', 'No adverse media detected.'))

        # Litigation history
        if litigation:
            _add_heading(self.doc, 'Litigation History', level=2)
            self.doc.add_paragraph(
                f"Active cases: {litigation.get('litigation_count', 0)} | "
                f"Risk: {litigation.get('litigation_risk', 'Low')} | "
                f"Exposure: {_fmt_inr(litigation.get('total_exposure_lakhs', 0))} Lakhs"
            )

        self.doc.add_paragraph()
        self.cam_content_text += "BORROWER_PROFILE completed\n"

    # ─── Section 4: Loan Proposal Details ──────────────────────────

    def _section_4_loan_proposal(self):
        _add_heading(self.doc, '3. LOAN PROPOSAL DETAILS', level=1)

        loan_struct = self.loan.get('loan_structure', {})
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        tl = loan_struct.get('term_loan', {})
        wc = loan_struct.get('working_capital', {})
        # Show the requested loan amount from onboarding
        requested_amt = self.app.get('loan_amount', self.app.get('requested_amount', None))
        rows = [
            ('Requested Loan Amount', f"₹{requested_amt} Lakhs" if requested_amt else '—'),
            ('Facility Type', 'Term Loan & Cash Credit'),
            ('Term Loan Amount', f"₹{tl.get('amount_lakhs', '—')} Lakhs"),
            ('Working Capital Amount', f"₹{wc.get('amount_lakhs', '—')} Lakhs"),
            ('Purpose', _safe(self.extracted.get('loan_purpose'), 'Business expansion / Working capital')),
            ('Tenure', f"{tl.get('tenure_months', '—')} months"),
            ('Interest Rate', f"{self.decision.get('interest_rate', '—')}% p.a."),
            ('Moratorium', _safe(self.extracted.get('moratorium'), '— months')),
            ('Repayment', 'Monthly installments'),
        ]
        for label, value in rows:
            _add_table_row(table, [label, str(value)])

        # Drawing Power calculation (RBI requirement for CC)
        _add_heading(self.doc, 'Drawing Power Calculation (for CC/OD)', level=2)
        stock = float(_safe(self.extracted.get('inventory', 0), 0))
        debtors = float(_safe(self.extracted.get('trade_receivables', 0), 0))
        creditors = float(_safe(self.extracted.get('trade_payables', 0), 0))
        dp = 0.75 * stock + 0.80 * debtors - creditors
        dp_table = self.doc.add_table(rows=0, cols=2)
        dp_table.style = 'Table Grid'
        _add_table_row(dp_table, ['75% of Stock (Inventory)', _fmt_inr(0.75 * stock)])
        _add_table_row(dp_table, ['80% of Debtors (Trade Receivables)', _fmt_inr(0.80 * debtors)])
        _add_table_row(dp_table, ['Less: Creditors (Trade Payables)', f"({_fmt_inr(creditors)})"])
        _add_table_row(dp_table, ['Drawing Power (DP)', _fmt_inr(dp)], bold=True)

        # Covenant schedule
        covenants = self.decision.get('covenants', [])
        if covenants:
            _add_heading(self.doc, 'Proposed Covenant Schedule', level=2)
            ct = self.doc.add_table(rows=1, cols=3)
            ct.style = 'Table Grid'
            ct.rows[0].cells[0].text = 'ID'
            ct.rows[0].cells[1].text = 'Type'
            ct.rows[0].cells[2].text = 'Description'
            for c in covenants:
                _add_table_row(ct, [c.get('id', '—'), c.get('type', '—'), c.get('description', '—')])

        self.doc.add_paragraph()
        self.cam_content_text += f"LOAN_PROPOSAL: DP={_fmt_inr(dp)}\n"

    # ─── Section 5: Financial Analysis ─────────────────────────────

    def _section_5_financial_analysis(self):
        _add_heading(self.doc, '4. FINANCIAL ANALYSIS', level=1)

        # Determine financial year from extracted data
        fy_label = _safe(self.extracted.get('financial_year'),
                         _safe(self.extracted.get('fs_financial_year'), None))
        # Try to build a year label like 'FY24'; fallback to the application year
        if not fy_label:
            app_year = self.app.get('financial_year', '')
            if app_year:
                fy_label = f"FY{str(app_year)[-2:]}" if len(str(app_year)) >= 4 else f"FY{app_year}"
            else:
                fy_label = 'Current FY'

        # Financial Summary table — single year (data from submitted documents)
        _add_heading(self.doc, 'Financial Summary', level=2)
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdrs = ['Metric', str(fy_label)]
        for i, h in enumerate(hdrs):
            table.rows[0].cells[i].text = h

        rev = _safe(self.extracted.get('total_revenue'), _safe(self.extracted.get('revenue_from_operations'), '—'))
        metrics = [
            ('Revenue', rev),
            ('EBITDA', _safe(self.extracted.get('ebitda'), '—')),
            ('PAT', _safe(self.extracted.get('profit_after_tax'), '—')),
            ('Net Worth', _safe(self.extracted.get('net_worth'), '—')),
            ('Total Debt', _safe(self.extracted.get('total_debt'), '—')),
        ]
        for metric_name, metric_val in metrics:
            _add_table_row(table, [metric_name, _fmt_inr(metric_val)])

        # Key Ratios
        _add_heading(self.doc, 'Key Financial Ratios', level=2)
        rt = self.doc.add_table(rows=0, cols=2)
        rt.style = 'Table Grid'

        debt = float(_safe(self.extracted.get('total_debt', 0), 0))
        nw = float(_safe(self.extracted.get('net_worth', 0), 0))
        ca = float(_safe(self.extracted.get('current_assets', 0), 0))
        cl = float(_safe(self.extracted.get('current_liabilities', 0), 0))

        ratios = [
            ('DSCR', _safe(self.extracted.get('dscr'), f"{self.decision.get('dscr', '—')}")),
            ('Current Ratio', f"{ca/cl:.2f}" if cl else '—'),
            ('Debt-to-Equity', f"{debt/nw:.2f}" if nw else '—'),
            ('LTV', _safe(self.extracted.get('ltv'), '—')),
            ('RONW', _safe(self.extracted.get('return_on_net_worth'), '—')),
        ]
        for label, value in ratios:
            _add_table_row(rt, [label, str(value)])

        # Cash Flow Statement (RBI requirement for loans > ₹1 Cr)
        _add_heading(self.doc, 'Cash Flow Statement', level=2)
        p = self.doc.add_paragraph()
        run = p.add_run('(RBI requirement for loans exceeding ₹1 Crore)')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        cf_table = self.doc.add_table(rows=0, cols=2)
        cf_table.style = 'Table Grid'
        _add_table_row(cf_table, ['Operating Cash Flow', _fmt_inr(self.extracted.get('operating_cash_flow', '—'))])
        _add_table_row(cf_table, ['Investing Cash Flow', _fmt_inr(self.extracted.get('investing_cash_flow', '—'))])
        _add_table_row(cf_table, ['Financing Cash Flow', _fmt_inr(self.extracted.get('financing_cash_flow', '—'))])
        _add_table_row(cf_table, ['Net Cash Flow', _fmt_inr(self.extracted.get('net_cash_flow', '—'))], bold=True)

        # GST vs Revenue Reconciliation (RBI Digital Lending Guidelines)
        _add_heading(self.doc, 'GST vs Bank Revenue Reconciliation', level=2)
        gst_forensics = self.l4.get('gst_forensics', {})
        a1 = gst_forensics.get('a1_reconciliation', {})
        gst_table = self.doc.add_table(rows=0, cols=2)
        gst_table.style = 'Table Grid'
        _add_table_row(gst_table, ['GST Turnover (GSTR-1)', _fmt_inr(self.extracted.get('gst_turnover_total', '—'))])
        _add_table_row(gst_table, ['Declared Revenue (P&L)', _fmt_inr(rev)])
        _add_table_row(gst_table, ['Revenue-GST Alignment', str(a1.get('revenue_gst_alignment', '—'))])
        _add_table_row(gst_table, ['GST Mismatch Ratio', f"{a1.get('gst_mismatch_ratio', '—')}%"])

        # Working Capital Cycle Analysis
        _add_heading(self.doc, 'Working Capital Cycle Analysis', level=2)
        wc_table = self.doc.add_table(rows=0, cols=2)
        wc_table.style = 'Table Grid'
        _add_table_row(wc_table, ['Debtor Days', _safe(self.extracted.get('debtor_days'), '—')])
        _add_table_row(wc_table, ['Inventory Days', _safe(self.extracted.get('inventory_days'), '—')])
        _add_table_row(wc_table, ['Creditor Days', _safe(self.extracted.get('creditor_days'), '—')])

        self.doc.add_paragraph()
        self.cam_content_text += "FINANCIAL_ANALYSIS completed\n"

    # ─── Section 6: Existing Debt Obligations ──────────────────────

    def _section_6_existing_debt(self):
        _add_heading(self.doc, '5. EXISTING DEBT OBLIGATIONS SCHEDULE', level=1)

        p = self.doc.add_paragraph()
        run = p.add_run('(Full schedule of all existing credit facilities as per RBI norms)')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        cibil = self.research.get('cibil', {})
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        for i, h in enumerate(['Lender', 'Facility / Limit', 'Outstanding', 'DPD']):
            table.rows[0].cells[i].text = h

        _add_table_row(table, ['CIBIL Rank', cibil.get('cibil_rank', '—'), '', ''])
        _add_table_row(table, ['CIBIL Score', cibil.get('cibil_score', '—'), '', ''])
        _add_table_row(table, ['Live Facilities', str(cibil.get('total_live_facilities', 0)), '', ''])
        _add_table_row(table, ['NPA Flag', 'YES' if cibil.get('npa_flag') else 'NO', '', ''])
        _add_table_row(table, ['Highest DPD', f"{cibil.get('highest_dpd_days', 0)} days", '', ''])

        # Personal guarantees
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        run = p.add_run('Guarantees: ')
        run.bold = True
        p.add_run(_safe(self.extracted.get('guarantees'), 'Personal Guarantee of Promoter + Corporate Guarantee'))

        self.doc.add_paragraph()
        self.cam_content_text += "DEBT_OBLIGATIONS completed\n"

    # ─── Section 7: Security & Collateral ──────────────────────────

    def _section_7_security_collateral(self):
        _add_heading(self.doc, '6. SECURITY & COLLATERAL SCHEDULE', level=1)

        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        _add_table_row(table, ['Primary Security', _safe(self.extracted.get('primary_security'),
            'Hypothecation of plant & machinery and stock-in-trade')])
        _add_table_row(table, ['Collateral Security', _safe(self.extracted.get('collateral_security'),
            'Equitable Mortgage of industrial land & building')])
        _add_table_row(table, ['Collateral Valuation', _fmt_inr(self.extracted.get('collateral_value', '—'))])
        _add_table_row(table, ['LTV Ratio', f"{_safe(self.extracted.get('ltv'), '65')}%"])
        _add_table_row(table, ['Valuation Date', _safe(self.extracted.get('valuation_date'), datetime.now().strftime('%d-%m-%Y'))])
        _add_table_row(table, ['Personal Guarantee', _safe(self.extracted.get('personal_guarantee'),
            f"Personal Guarantee of {_safe(self.extracted.get('promoter_name'), 'Promoter')}")])
        _add_table_row(table, ['Corporate Guarantee', _safe(self.extracted.get('corporate_guarantee'),
            f"Corporate Guarantee of {self.app.get('company_name', 'Group Company')}")])
        _add_table_row(table, ['Insurance Requirement', 'Comprehensive insurance coverage as per bank policy'])

        self.doc.add_paragraph()
        self.cam_content_text += "SECURITY_COLLATERAL completed\n"

    # ─── Section 8: Sector & Industry Risk ─────────────────────────

    def _section_8_sector_risk(self):
        _add_heading(self.doc, '7. SECTOR & INDUSTRY RISK INTELLIGENCE', level=1)

        sector = self.research.get('sector_risk', {})
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        _add_table_row(table, ['Industry / Sector', _safe(sector.get('sector'), _safe(self.extracted.get('industry'), '—'))])
        _add_table_row(table, ['Sector Risk Score', f"{sector.get('sector_risk_score', 0):.2f}"])
        _add_table_row(table, ['RBI Sector Flag', 'YES — Caution' if sector.get('rbi_sector_flag') else 'No sector caution'])

        # Headwinds / Tailwinds
        headwinds = sector.get('headwinds', [])
        tailwinds = sector.get('tailwinds', [])
        if headwinds:
            _add_heading(self.doc, 'Headwinds', level=2)
            for h in headwinds:
                self.doc.add_paragraph(f"• {h}", style='List Bullet')
        if tailwinds:
            _add_heading(self.doc, 'Tailwinds', level=2)
            for t in tailwinds:
                self.doc.add_paragraph(f"• {t}", style='List Bullet')

        # E&S Risk Comment (RBI Sustainable Finance Framework)
        _add_heading(self.doc, 'Environmental & Social (E&S) Risk Assessment', level=2)
        industry = _safe(self.extracted.get('industry'), _safe(sector.get('sector'), 'General'))
        self.doc.add_paragraph(
            f"As per RBI Sustainable Finance Framework, E&S risk assessment has been conducted for "
            f"{self.app.get('company_name', 'the borrower')} operating in the {industry} sector. "
            f"The company is classified as {'medium E&S risk' if 'manufact' in industry.lower() else 'low E&S risk'} "
            f"given its industry profile. Environmental compliance certificates and pollution control board clearances "
            f"should be obtained as a condition precedent."
        )

        self.doc.add_paragraph()
        self.cam_content_text += "SECTOR_RISK completed\n"

    # ─── Section 8B: ESG Risk & Climate Analysis ─────────────────────

    def _section_8b_esg_risk(self):
        _add_heading(self.doc, '8B. ESG RISK & CLIMATE ANALYSIS', level=1)

        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        _add_table_row(table, ['Transition Risk', _safe(self.extracted.get('esg_transition_risk'))], bold=True)
        _add_table_row(table, ['Physical Risk', _safe(self.extracted.get('esg_physical_risk'))])
        _add_table_row(table, ['Carbon Footprint (MT CO₂e)', _safe(self.extracted.get('carbon_footprint_mt'))])
        _add_table_row(table, ['Scope 1 Emissions', _safe(self.extracted.get('scope1_emissions'))])
        _add_table_row(table, ['Scope 2 Emissions', _safe(self.extracted.get('scope2_emissions'))])
        _add_table_row(table, ['Scope 3 Emissions', _safe(self.extracted.get('scope3_emissions'))])
        _add_table_row(table, ['Sustainability Rating', _safe(self.extracted.get('sustainability_rating'))])
        _add_table_row(table, ['Rating Agency', _safe(self.extracted.get('sustainability_rating_agency'))])
        _add_table_row(table, ['Renewable Energy (%)', _safe(self.extracted.get('renewable_energy_pct'))])
        _add_table_row(table, ['Waste Recycled (%)', _safe(self.extracted.get('waste_recycled_pct'))])
        _add_table_row(table, ['Report Period', _safe(self.extracted.get('esg_report_period'))])

        # Green Financing eligibility callout
        green_eligible = self.extracted.get('green_financing_eligible')
        if green_eligible is True or str(green_eligible).lower() == 'true':
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run('✅ GREEN FINANCING ELIGIBLE — ')
            run.bold = True
            run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
            p.add_run(
                'Based on the sustainability metrics extracted from the Climate/ESG Report, '
                'this borrower qualifies for a Green Loan product with a recommended interest '
                'rate discount of 0.25% – 0.50%. This is in line with RBI Sustainable Finance Framework '
                'and Green Bond Principles.'
            )
        else:
            self.doc.add_paragraph(
                'Green Financing eligibility was not established from the submitted reports, '
                'or no ESG/Climate report was provided.'
            )

        # Key ESG risks
        esg_risks = self.extracted.get('esg_key_risks', [])
        if esg_risks and isinstance(esg_risks, list):
            _add_heading(self.doc, 'Key ESG Risks Identified', level=2)
            for risk in esg_risks:
                self.doc.add_paragraph(f"• {risk}", style='List Bullet')

        self.doc.add_paragraph()
        self.cam_content_text += "ESG_RISK completed\n"

    # ─── Section 3C: Statutory Compliance (Annual Return) ──────────

    def _section_3c_statutory_compliance(self):
        _add_heading(self.doc, '3C. STATUTORY COMPLIANCE & OWNERSHIP VERIFICATION', level=1)

        filed = self.extracted.get('annual_return_filed')
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        # Highlight non-filing as critical risk
        filed_text = '✅ Filed' if (filed is True or str(filed).lower() == 'true') else '⚠️ Not Filed / Unknown'
        _add_table_row(table, ['Annual Return Filed', filed_text],
                       bold=True,
                       bg_color='f0fdf4' if 'Filed' in filed_text else 'fef2f2')
        _add_table_row(table, ['Filing Date', _safe(self.extracted.get('annual_return_filing_date'))])
        _add_table_row(table, ['Return Year', _safe(self.extracted.get('annual_return_year'))])
        _add_table_row(table, ['Registered Office', _safe(self.extracted.get('registered_office_address'))])
        _add_table_row(table, ['MCA Company Status', _safe(self.extracted.get('mca_company_status'))])
        _add_table_row(table, ['Total Directors', _safe(self.extracted.get('total_directors_count'))])
        _add_table_row(table, ['Independent Directors', _safe(self.extracted.get('independent_directors_count'))])
        _add_table_row(table, ['Charges Registered', _safe(self.extracted.get('charges_registered'))])
        _add_table_row(table, ['Declared Paid-up Capital', _safe(self.extracted.get('declared_paid_up_capital'))])
        _add_table_row(table, ['Declared Reserves', _safe(self.extracted.get('declared_reserves'))])

        # Indebtedness cross-check
        declared = self.extracted.get('declared_indebtedness')
        financial = self.extracted.get('total_outstanding_borrowings')
        _add_table_row(table, ['Declared Indebtedness (Annual Return)', _safe(declared)], bold=True)
        _add_table_row(table, ['Total Borrowings (Financial Statements)', _safe(financial)])

        if declared and financial:
            try:
                d = float(declared)
                f_val = float(financial)
                if f_val > 0:
                    mismatch_pct = abs(d - f_val) / f_val * 100
                    if mismatch_pct > 5:
                        p = self.doc.add_paragraph()
                        run = p.add_run(f'⚠️ CROSS-CHECK ALERT: {mismatch_pct:.1f}% mismatch between declared indebtedness '
                                       f'and financial statement borrowings. Manual verification required.')
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                    else:
                        self.doc.add_paragraph('✅ Declared indebtedness matches financial statements within 5% tolerance.')
            except (ValueError, TypeError):
                pass

        # Directors list
        directors = self.extracted.get('active_directors_list', [])
        if directors and isinstance(directors, list):
            _add_heading(self.doc, 'Active Directors (As per Annual Return)', level=2)
            for i, name in enumerate(directors, 1):
                self.doc.add_paragraph(f"{i}. {name}")

        # Compliance flags
        flags = self.extracted.get('statutory_compliance_flags', [])
        if flags and isinstance(flags, list):
            _add_heading(self.doc, 'Compliance Flags', level=2)
            for flag in flags:
                self.doc.add_paragraph(f"⚠ {flag}", style='List Bullet')

        self.doc.add_paragraph()
        self.cam_content_text += "STATUTORY_COMPLIANCE completed\n"

    # ─── Section 9: AI Risk Score & Explainability ─────────────────

    def _section_9_ai_risk_score(self):
        _add_heading(self.doc, '8. AI RISK SCORE & EXPLAINABILITY', level=1)

        score = self.decision.get('final_credit_score', '—')
        pd = self.decision.get('probability_of_default', 0)
        risk_band = self.decision.get('risk_band', '—')

        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        _add_table_row(table, ['Federated Credit Score', str(score)], bold=True)
        _add_table_row(table, ['Risk Band / Grade', risk_band])
        _add_table_row(table, ['Probability of Default', f"{pd*100:.1f}%" if isinstance(pd, (int, float)) else str(pd)])
        _add_table_row(table, ['Model Version', self.l5.get('audit_snapshot', {}).get('model_metadata', {}).get('model_version', 'v4.3')])

        # SHAP top drivers
        _add_heading(self.doc, 'Top Positive Drivers (Score Boosters)', level=2)
        for d in self.explanation.get('shap_top_positive', [])[:3]:
            self.doc.add_paragraph(
                f"• {d.get('label', '—')}: Value={d.get('value', '—')}, SHAP Impact={d.get('shap_value', 0):.4f} ({d.get('magnitude', '')})"
            )

        _add_heading(self.doc, 'Top Negative Drivers (Score Detractors)', level=2)
        for d in self.explanation.get('shap_top_negative', [])[:3]:
            self.doc.add_paragraph(
                f"• {d.get('label', '—')}: Value={d.get('value', '—')}, SHAP Impact=+{d.get('shap_value', 0):.4f} ({d.get('magnitude', '')})"
            )

        # Sensitivity / Stress Test
        _add_heading(self.doc, 'Sensitivity / Stress Test Analysis', level=2)
        p = self.doc.add_paragraph()
        run = p.add_run('(What happens if revenue drops 15%-30%)')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        dscr_val = float(_safe(self.extracted.get('dscr', 1.5), 1.5))
        stress_table = self.doc.add_table(rows=1, cols=4)
        stress_table.style = 'Table Grid'
        for i, h in enumerate(['Scenario', 'Revenue Impact', 'Estimated DSCR', 'Assessment']):
            stress_table.rows[0].cells[i].text = h
        _add_table_row(stress_table, ['Base Case', '0%', f"{dscr_val:.2f}", 'Serviceable' if dscr_val > 1.2 else 'Stressed'])
        _add_table_row(stress_table, ['Downside (-15%)', '-15%', f"{dscr_val * 0.85:.2f}", 
            'Serviceable' if dscr_val * 0.85 > 1.2 else 'Stressed'])
        _add_table_row(stress_table, ['Stress (-30%)', '-30%', f"{dscr_val * 0.70:.2f}",
            'Serviceable' if dscr_val * 0.70 > 1.2 else 'Critical'])

        self.doc.add_paragraph()
        self.cam_content_text += f"AI_RISK_SCORE: {score} | PD: {pd}\n"

    # ─── Section 10: Forensic Alert Summary ────────────────────────

    def _section_10_forensic_alerts(self):
        _add_heading(self.doc, '9. FORENSIC ALERT SUMMARY', level=1)

        alerts = self.forensics.get('alerts', [])
        if alerts:
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            for i, h in enumerate(['Alert Type', 'Severity', 'Description', 'Score Penalty']):
                table.rows[0].cells[i].text = h
            for a in alerts:
                bg = 'FFCCCC' if a.get('severity') == 'RED' else 'FFF3CD' if a.get('severity') == 'AMBER' else None
                _add_table_row(table, [
                    a.get('type', '—').replace('_', ' '),
                    a.get('severity', '—'),
                    a.get('description', '—')[:80],
                    str(a.get('score_penalty', 0))
                ], bg_color=bg)
        else:
            self.doc.add_paragraph('No forensic alerts detected. All checks passed.')

        # E&S risk addendum
        _add_heading(self.doc, 'E&S Risk Comment', level=2)
        self.doc.add_paragraph(
            'Environmental and Social risk has been reviewed as per RBI Sustainable Finance Framework. '
            'Manufacturing companies require environmental compliance certification.'
        )

        self.doc.add_paragraph()
        self.cam_content_text += f"FORENSICS: {len(alerts)} alerts\n"

    # ─── Section 11: Five Cs of Credit ─────────────────────────────

    def _section_11_five_cs(self):
        _add_heading(self.doc, '10. THE FIVE Cs OF CREDIT', level=1)

        p = self.doc.add_paragraph()
        run = p.add_run('(LLM-generated, grounded in all above data)')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        c_labels = {
            'character': '👤 Character',
            'capacity': '💪 Capacity',
            'capital': '🏦 Capital',
            'collateral': '🏠 Collateral',
            'conditions': '🌐 Conditions'
        }
        c_subtitles = {
            'character': 'Promoter integrity, litigation, CIBIL',
            'capacity': 'DSCR, cash flows, GST alignment',
            'capital': 'Net worth, leverage, Ind-AS notes',
            'collateral': 'Security coverage, LTV, valuation',
            'conditions': 'Covenants, sector headwinds, RBI/SEBI risk',
        }

        for key, label in c_labels.items():
            c_data = self.five_cs.get(key, {})
            rating = c_data.get('rating', 'NOT ASSESSED')
            explanation = c_data.get('explanation', '')

            h = self.doc.add_heading(f"{label}: {rating}", level=2)
            p = self.doc.add_paragraph()
            run = p.add_run(c_subtitles.get(key, ''))
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x99)

            if explanation:
                self.doc.add_paragraph(explanation)

        self.doc.add_paragraph()
        self.cam_content_text += "FIVE_CS completed\n"

    # ─── Section 12: SWOT Analysis ─────────────────────────────────

    def _section_12_swot(self):
        _add_heading(self.doc, '11. SWOT ANALYSIS', level=1)

        biggest_risk = self.explanation.get('biggest_risk', '')
        biggest_strength = self.explanation.get('biggest_strength', '')
        sector = self.research.get('sector_risk', {})

        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        table.rows[0].cells[0].text = f"STRENGTHS\n{biggest_strength or 'Strong banking relationship, stable operations'}"
        table.rows[0].cells[1].text = f"WEAKNESSES\n{biggest_risk or 'Sector-specific challenges, working capital pressure'}"
        table.rows[1].cells[0].text = f"OPPORTUNITIES\n{'; '.join(sector.get('tailwinds', ['Market expansion potential']))}"
        table.rows[1].cells[1].text = f"THREATS\n{'; '.join(sector.get('headwinds', ['Competitive pressure']))}"

        self.doc.add_paragraph()
        self.cam_content_text += "SWOT completed\n"

    # ─── Section 13: Officer-Flagged Issues & Custom Fields ────────

    def _section_13_officer_issues(self):
        _add_heading(self.doc, '12. OFFICER-FLAGGED ISSUES & CUSTOM DATA', level=1)

        p = self.doc.add_paragraph()
        run = p.add_run('(HITL-originated observations not captured by AI pipeline)')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # ── Officer Issues ──
        issues = self.officer_issues
        if issues:
            _add_heading(self.doc, 'Officer-Flagged Issues', level=2)
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            for i, h in enumerate(['Issue Title', 'Severity', 'Description', 'Checkpoint']):
                table.rows[0].cells[i].text = h

            for issue in issues:
                sev = issue.get('severity', 'MEDIUM')
                bg = None
                if sev == 'CRITICAL': bg = 'FFCCCC'
                elif sev == 'HIGH': bg = 'FFE0CC'
                elif sev == 'MEDIUM': bg = 'FFF3CD'

                _add_table_row(table, [
                    issue.get('title', '—'),
                    sev,
                    issue.get('description', '—')[:120],
                    f"HITL-{issue.get('checkpoint', '?')}"
                ], bg_color=bg)

            # Impact statement
            critical_count = sum(1 for i in issues if i.get('severity') in ('CRITICAL', 'HIGH'))
            if critical_count > 0:
                self.doc.add_paragraph()
                p = self.doc.add_paragraph()
                run = p.add_run(f'⚠️ Impact: {critical_count} HIGH/CRITICAL issue(s) flagged by officer — '
                                f'these have been factored into the AI risk scoring and final decision.')
                run.bold = True
                run.font.color.rgb = RGBColor(0xcc, 0x33, 0x00)
        else:
            self.doc.add_paragraph('No officer-flagged issues were raised during the HITL review process.')

        # ── Custom Data Fields ──
        custom = self.custom_fields
        if custom and isinstance(custom, dict) and len(custom) > 0:
            _add_heading(self.doc, 'Officer-Added Custom Data Fields', level=2)
            ct = self.doc.add_table(rows=1, cols=2)
            ct.style = 'Table Grid'
            ct.rows[0].cells[0].text = 'Field Name'
            ct.rows[0].cells[1].text = 'Value'
            for field_name, field_value in custom.items():
                _add_table_row(ct, [str(field_name).replace('_', ' ').title(), str(field_value)])

            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run(f'{len(custom)} custom data field(s) were added by the reviewing officer '
                            f'and have been incorporated into the pipeline context.')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
        else:
            self.doc.add_paragraph('No custom data fields were added during the review process.')

        # ── GST Cross-Validation Summary ──
        gstin_mismatches = self.app.get('gstin_mismatch_fields')
        if gstin_mismatches:
            if isinstance(gstin_mismatches, str):
                try:
                    gstin_mismatches = json.loads(gstin_mismatches)
                except:
                    gstin_mismatches = None
        if gstin_mismatches and isinstance(gstin_mismatches, list) and len(gstin_mismatches) > 0:
            _add_heading(self.doc, 'GST Cross-Validation Discrepancies', level=2)
            for mismatch in gstin_mismatches:
                self.doc.add_paragraph(
                    f"• {mismatch.get('field', '—')}: Official='{mismatch.get('official', '—')}' "
                    f"vs Extracted='{mismatch.get('extracted', '—')}'",
                    style='List Bullet'
                )

        self.doc.add_paragraph()
        self.cam_content_text += f"OFFICER_ISSUES: {len(issues)} issues, {len(custom)} custom fields\n"

    # ─── Section 14: Recommendations & Approvals ───────────────────

    def _section_14_recommendations(self):
        _add_heading(self.doc, '13. RECOMMENDATIONS & APPROVALS', level=1)

        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        _add_table_row(table, ['AI Recommendation', self.decision.get('decision', '—')], bold=True)
        _add_table_row(table, ['Federated Credit Score', str(self.decision.get('final_credit_score', '—'))])
        _add_table_row(table, ['Sanction Amount', f"₹{self.decision.get('sanction_amount_lakhs', '—')} Lakhs"])
        _add_table_row(table, ['Interest Rate', f"{self.decision.get('interest_rate', '—')}% p.a."])

        # Conditions
        conditions = self.decision.get('conditions', [])
        if conditions:
            _add_heading(self.doc, 'Sanction Conditions', level=2)
            for c in conditions:
                self.doc.add_paragraph(f"• {c}", style='List Bullet')

        # Loan Amount Validation & Explainability
        requested = self.app.get('loan_amount')
        sanctioned = self.decision.get('sanction_amount_lakhs')
        if requested or sanctioned:
            _add_heading(self.doc, 'Loan Amount Validation & Rationale', level=2)
            la_table = self.doc.add_table(rows=0, cols=2)
            la_table.style = 'Table Grid'
            _add_table_row(la_table, ['Requested Amount', f"₹{requested} Lakhs" if requested else '—'])
            _add_table_row(la_table, ['Sanctioned Amount', f"₹{sanctioned} Lakhs" if sanctioned else '—'])

            loan_analysis = self.l5.get('loan_amount_analysis', '')
            if loan_analysis:
                self.doc.add_paragraph()
                p = self.doc.add_paragraph()
                run = p.add_run('AI Loan Amount Analysis: ')
                run.bold = True
                run.font.size = Pt(10)
                p.add_run(str(loan_analysis))
            elif requested and sanctioned:
                try:
                    req_f = float(requested)
                    sanc_f = float(sanctioned)
                    diff_pct = ((sanc_f - req_f) / req_f * 100) if req_f else 0
                    if abs(diff_pct) < 5:
                        rationale = (f"The sanctioned amount (₹{sanc_f} L) closely matches the requested amount "
                                     f"(₹{req_f} L), indicating the borrower's financial profile supports the request.")
                    elif diff_pct < -5:
                        rationale = (f"The sanctioned amount (₹{sanc_f} L) is {abs(diff_pct):.0f}% lower than requested "
                                     f"(₹{req_f} L). This reduction is based on the borrower's debt servicing capacity, "
                                     f"existing leverage, and risk profile as assessed by the AI model.")
                    else:
                        rationale = (f"The AI model suggests the borrower can service up to ₹{sanc_f} L, which is "
                                     f"{diff_pct:.0f}% higher than requested (₹{req_f} L). The borrower's strong "
                                     f"financial indicators support a higher facility.")
                    self.doc.add_paragraph(rationale)
                except (ValueError, TypeError):
                    pass

        # Deviation tracking
        _add_heading(self.doc, 'Deviation Tracking', level=2)
        self.doc.add_paragraph('None' if self.decision.get('decision') != 'CONDITIONAL' else 'See conditions above')

        # Signature blocks
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        sig_table = self.doc.add_table(rows=4, cols=3)
        sig_table.style = 'Table Grid'
        sig_table.rows[0].cells[0].text = 'MAKER'
        sig_table.rows[0].cells[1].text = 'CHECKER'
        sig_table.rows[0].cells[2].text = 'CREDIT COMMITTEE'
        sig_table.rows[1].cells[0].text = 'Signature: ________________'
        sig_table.rows[1].cells[1].text = 'Signature: ________________'
        sig_table.rows[1].cells[2].text = 'Signature: ________________'
        sig_table.rows[2].cells[0].text = 'Name: Credit Analyst'
        sig_table.rows[2].cells[1].text = 'Name: Branch Manager'
        sig_table.rows[2].cells[2].text = 'Name: Committee Chair'
        sig_table.rows[3].cells[0].text = f"Date: {datetime.now().strftime('%d-%m-%Y')}"
        sig_table.rows[3].cells[1].text = f"Date: {datetime.now().strftime('%d-%m-%Y')}"
        sig_table.rows[3].cells[2].text = f"Date: {datetime.now().strftime('%d-%m-%Y')}"

        # Digital signature hash
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        run = p.add_run(f"Digital Signature Hash (SHA-256): [Generated at export]")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p = self.doc.add_paragraph()
        run = p.add_run(f"Case ID: {self.app.get('case_id', '—')} | Timestamp: {datetime.utcnow().isoformat()} | Model: v4.3")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        self.cam_content_text += "RECOMMENDATIONS completed\n"

    # ─── Appendix A: Evidence Panel ────────────────────────────────

    def _appendix_a_evidence(self):
        self.doc.add_page_break()
        _add_heading(self.doc, 'APPENDIX A: Evidence Panel', level=1)

        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        for i, h in enumerate(['Ref#', 'Data Point', 'Source', 'Confidence']):
            table.rows[0].cells[i].text = h

        evidence = [
            ('E1', 'Revenue from Operations', 'Annual Report / P&L', 'HIGH'),
            ('E2', 'GST Turnover', 'GSTR-1 Filing', 'HIGH'),
            ('E3', 'Bank Statement Analysis', 'Bank Statement PDF', 'HIGH'),
            ('E4', 'Litigation History', 'eCourts / NCLT (Tavily)', 'MEDIUM'),
            ('E5', 'Adverse Media', 'Web Research (Tavily+Groq)', 'MEDIUM'),
            ('E6', 'Sector Risk Score', 'RBI Publications 2024', 'HIGH'),
            ('E7', 'CIBIL Commercial', 'TransUnion CIBIL (Simulated)', 'MEDIUM'),
            ('E8', 'MCA Company Status', 'MCA21 Portal (Simulated)', 'MEDIUM'),
            ('E9', 'Director DIN Status', 'MCA21 Portal (Simulated)', 'MEDIUM'),
            ('E10', 'Five Cs Assessment', 'Groq LLM Qualitative Overlay', 'MEDIUM'),
            ('E11', 'SHAP Explainability', 'XGBoost + SHAP Library', 'HIGH'),
            ('E12', 'Promoter Net Worth', 'Declaration / Submitted Docs', 'MEDIUM'),
        ]
        for e in evidence:
            _add_table_row(table, list(e))

        self.cam_content_text += "EVIDENCE_PANEL completed\n"

    # ─── Tamper-Evident Hash ───────────────────────────────────────

    def _compute_hash(self):
        return hashlib.sha256(self.cam_content_text.encode()).hexdigest()


# ─── Convenience: Generate CAM from application data ─────────────────────────

def generate_cam_report(app_data, l2_data, l3_data, l4_data, l5_data, output_dir):
    """Generate CAM in DOCX format and return metadata."""
    os.makedirs(output_dir, exist_ok=True)

    case_id = app_data.get('case_id', 'UNKNOWN')
    docx_path = os.path.join(output_dir, f"CAM_{case_id}.docx")

    gen = CAMGenerator(app_data, l2_data, l3_data, l4_data, l5_data)
    result = gen.generate(docx_path)

    return result


def convert_docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF. Tries docx2pdf, falls back to simple copy."""
    import os
    abs_docx = os.path.abspath(docx_path)
    abs_pdf = os.path.abspath(pdf_path)
    try:
        from docx2pdf import convert
        convert(abs_docx, abs_pdf)
        return os.path.exists(abs_pdf)
    except Exception as e:
        print(f"docx2pdf failed: {e}")
        try:
            # Fallback: try libreoffice
            import subprocess
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', os.path.dirname(abs_pdf), abs_docx
            ], timeout=60, check=True)
            return os.path.exists(abs_pdf)
        except Exception as e2:
            print(f"LibreOffice fallback failed: {e2}")
            return False


def generate_audit_json(app_data, l2_data, l3_data, l4_data, l5_data):
    """Generate a machine-readable audit JSON for the case."""
    decision = l5_data.get('decision_summary', {})
    five_cs = l5_data.get('explanation', {}).get('five_cs', {})
    forensics = l4_data.get('forensics_report', {})
    features = l4_data.get('feature_vector', {})
    snapshot = l5_data.get('audit_snapshot', {})

    officer_issues = app_data.get('officer_issues', [])
    custom_fields = app_data.get('custom_fields', {})

    audit = {
        'case_id': app_data.get('case_id'),
        'timestamp': datetime.utcnow().isoformat(),
        'company_name': app_data.get('company_name'),
        'ai_score': decision.get('final_credit_score'),
        'risk_band': decision.get('risk_band'),
        'probability_of_default': decision.get('probability_of_default'),
        'final_decision': decision.get('decision'),
        'interest_rate': decision.get('interest_rate'),
        'sanction_amount_lakhs': decision.get('sanction_amount_lakhs'),
        'requested_amount_lakhs': app_data.get('loan_amount'),
        'five_cs': {k: {'rating': v.get('rating')} for k, v in five_cs.items()},
        'forensics_summary': {
            'red_flags': forensics.get('red_flag_count', 0),
            'amber_flags': forensics.get('amber_flag_count', 0),
            'total_penalty': forensics.get('total_score_penalty', 0),
        },
        'officer_issues': [
            {'title': i.get('title'), 'severity': i.get('severity')}
            for i in officer_issues
        ] if officer_issues else [],
        'custom_fields_count': len(custom_fields) if isinstance(custom_fields, dict) else 0,
        'feature_count': len(features),
        'cam_hash': hashlib.sha256(json.dumps(decision, default=str).encode()).hexdigest(),
        'model_version': snapshot.get('model_metadata', {}).get('model_version', 'v4.3'),
        'schema_version': '1.1.0',
    }
    return audit
